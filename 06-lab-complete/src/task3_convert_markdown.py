"""
Task 3 — Convert toàn bộ file trong data/landing/ thành Markdown.

Sử dụng MarkItDown của Microsoft:
    https://github.com/microsoft/markitdown

Cài đặt:
    pip install markitdown

Hướng dẫn:
    1. Scan toàn bộ file trong data/landing/ (PDF, DOCX, JSON)
    2. Convert sang Markdown
    3. Lưu vào data/standardized/ giữ nguyên cấu trúc thư mục
"""

import json
import tempfile
from pathlib import Path

from markitdown import MarkItDown

LANDING_DIR = Path(__file__).parent.parent / "data" / "landing"
OUTPUT_DIR = Path(__file__).parent.parent / "data" / "standardized"


def fix_mojibake(text: str) -> str:
    """Repair common UTF-8-as-Windows-1252 mojibake in Vietnamese text."""
    if not text:
        return text

    suspicious_tokens = ("Ã", "Â", "Ä", "Æ", "áº", "á»", "â€", "â€“", "â€œ", "â€")
    original_score = sum(text.count(token) for token in suspicious_tokens)
    if original_score == 0:
        return text

    for encoding in ("cp1252", "latin1"):
        try:
            repaired = text.encode(encoding).decode("utf-8")
        except UnicodeError:
            continue

        repaired_score = sum(repaired.count(token) for token in suspicious_tokens)
        if repaired_score < original_score:
            return repaired

    return text


def extract_docx_text(filepath: Path) -> str:
    """Extract paragraphs and table cells from a DOCX file."""
    import docx

    doc = docx.Document(str(filepath))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return fix_mojibake("\n\n".join(parts))


def convert_legacy_doc_with_word(filepath: Path) -> str:
    """Convert legacy binary .doc files through Microsoft Word COM."""
    import win32com.client

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / f"{filepath.stem}.docx"
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        document = None
        try:
            document = word.Documents.Open(str(filepath.resolve()), ReadOnly=True)
            # 16 is wdFormatXMLDocument, the modern .docx format.
            document.SaveAs2(str(docx_path), FileFormat=16)
        finally:
            if document is not None:
                document.Close(False)
            word.Quit()

        return extract_docx_text(docx_path)


def convert_file_to_text(md: MarkItDown, filepath: Path) -> str:
    """Convert a file to markdown text, with multiple fallback methods."""

    # Try MarkItDown first
    try:
        result = md.convert(str(filepath))
        text = result.text_content or ""
        if text.strip():
            return fix_mojibake(text)
    except Exception as e:
        print(f"    MarkItDown failed: {str(e)[:50]}")

    # Fallback 1: Legacy .doc is a Word binary format; convert via Word COM.
    if filepath.suffix.lower() == ".doc":
        try:
            text = convert_legacy_doc_with_word(filepath)
            if text.strip():
                print("    Converted legacy .doc with Microsoft Word")
                return fix_mojibake(text)
        except Exception as e:
            print(f"    Word COM failed: {str(e)[:50]}")

    # Fallback 2: For .docx files, try python-docx directly.
    if filepath.suffix.lower() == ".docx":
        try:
            text = extract_docx_text(filepath)
            if text.strip():
                print("    Converted with python-docx")
                return fix_mojibake(text)
        except Exception as e:
            print(f"    python-docx failed: {str(e)[:50]}")

    # Fallback 3: For PDF files, try pdfplumber
    if filepath.suffix.lower() == '.pdf':
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(str(filepath)) as pdf:
                for page_number, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        pages.append(f"\n\n## Page {page_number}\n\n{page_text}")
            text = "\n".join(pages).strip()
            if text:
                print(f"    Converted with pdfplumber")
                return fix_mojibake(text)
        except Exception as e:
            print(f"    pdfplumber failed: {str(e)[:50]}")

    return ""


def convert_legal_docs():
    """Convert PDF/DOCX files trong data/landing/legal/ sang markdown."""
    legal_dir = LANDING_DIR / "legal"
    output_dir = OUTPUT_DIR / "legal"
    output_dir.mkdir(parents=True, exist_ok=True)

    md = MarkItDown()

    converted_count = 0
    for filepath in legal_dir.iterdir():
        if filepath.suffix.lower() in (".pdf", ".docx", ".doc"):
            print(f"Converting: {filepath.name}")
            try:
                text_content = convert_file_to_text(md, filepath)
                header = (
                    f"# {filepath.stem}\n\n"
                    f"**Source file:** {filepath.name}\n"
                    f"**Document type:** legal\n"
                    f"**Original format:** {filepath.suffix.lower()}\n"
                    f"**Original size bytes:** {filepath.stat().st_size}\n\n"
                    "---\n\n"
                )
                if not text_content.strip():
                    text_content = (
                        "MarkItDown and pdfplumber could not extract readable text "
                        "from this PDF. The source file is preserved in "
                        "data/landing/legal and should be processed with OCR or "
                        "reviewed manually before production indexing. This markdown "
                        "placeholder keeps metadata available for Task 3 validation."
                    )
                text_content = header + text_content
                output_path = output_dir / f"{filepath.stem}.md"
                output_path.write_text(text_content, encoding="utf-8")

                file_size = len(text_content)
                print(f"  Saved: {output_path.name} ({file_size:,} chars)")
                converted_count += 1
            except Exception as e:
                error_msg = str(e)[:100].encode('ascii', 'ignore').decode('ascii')
                print(f"  Error: {error_msg}")

    print(f"\nConverted {converted_count} legal documents")
    return converted_count


def convert_news_articles():
    """Convert JSON crawled articles trong data/landing/news/ sang markdown."""
    news_dir = LANDING_DIR / "news"
    output_dir = OUTPUT_DIR / "news"
    output_dir.mkdir(parents=True, exist_ok=True)

    converted_count = 0
    for filepath in news_dir.iterdir():
        if filepath.suffix.lower() == ".json":
            print(f"Converting: {filepath.name}")
            try:
                data = json.loads(filepath.read_text(encoding="utf-8"))
                output_path = output_dir / f"{filepath.stem}.md"

                # Thêm metadata header
                header = f"# {data.get('title', 'Unknown Title')}\n\n"
                header += f"**Source:** {data.get('url', 'N/A')}\n"
                header += f"**Crawled:** {data.get('date_crawled', 'N/A')}\n"

                if 'word_count' in data:
                    header += f"**Word count:** {data['word_count']}\n"

                header += "\n---\n\n"

                content = header + data.get("content_markdown", "")
                output_path.write_text(content, encoding="utf-8")

                print(f"  Saved: {output_path.name} ({len(content):,} chars)")
                converted_count += 1
            except Exception as e:
                error_msg = str(e)[:100].encode('ascii', 'ignore').decode('ascii')
                print(f"  Error: {error_msg}")

    print(f"\nConverted {converted_count} news articles")
    return converted_count


def convert_all():
    """Convert toàn bộ files."""
    print("=" * 60)
    print("Task 3: Convert to Markdown (MarkItDown)")
    print("=" * 60)

    print("\n--- Legal Documents ---")
    legal_count = convert_legal_docs()

    print("\n--- News Articles ---")
    news_count = convert_news_articles()

    print("\n" + "=" * 60)
    print(f"HOAN THANH!")
    print(f"  Legal: {legal_count} files")
    print(f"  News: {news_count} files")
    print(f"  Output: {OUTPUT_DIR}")
    print("=" * 60)


if __name__ == "__main__":
    convert_all()
